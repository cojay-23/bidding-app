from __future__ import annotations

import html
import json
import os
import re
import shutil
import subprocess
import sys
import zipfile
from collections import Counter
from datetime import date, datetime
from pathlib import Path
from typing import Any

from docx import Document
from pypdf import PdfReader

from .config import PROJECTS_DIR
from .database import add_report, list_files, stored_file_path, update_project


TODAY = date(2026, 6, 21)
# ═══════════════════════════════════════════════════════════════
# 废标条款关键词体系（参照 bidding-analyst skill 四级分类）
# ═══════════════════════════════════════════════════════════════
RISK_FATAL = (  # 🔴 致命 — 直接废标表述
    "废标", "否决", "作废", "不予通过", "取消中标资格",
    "按废标处理", "视为不响应", "无效投标", "投标无效",
    "拒绝", "取消投标资格",
)
RISK_STAR = (  # ⭐ 星号否决条款 — 一条不满足即废标
    "★", "*",
)
RISK_STRICT = (  # 🟡 严重 — 强制/禁止性要求
    "必须", "不得", "严禁", "禁止", "否则废标", "否则无效",
    "应", "须", "投标人应", "投标人须",
)
RISK_REVIEW = (  # 🟠 需注意 — 审查相关
    "资格审查", "符合性审查", "初步评审", "形式审查",
    "资格条件", "资质要求", "不接受", "不予受理",
    "逾期", "负偏离", "不满足",
)
KEY_RISK_WORDS = RISK_FATAL + RISK_STRICT + RISK_REVIEW

SCORE_WORDS = ("评分", "得分", "分值", "满分", "技术", "商务", "价格", "评标办法", "评分标准")
MATERIAL_WORDS = ("提供", "证明", "截图", "承诺函", "授权", "营业执照", "保函", "查询结果", "证书", "复印件", "扫描件")
SCREENSHOT_WORDS = ("截图", "查询截图", "功能截图", "页面截图", "网站截图", "平台截图", "系统截图", "查询结果")
QUALIFICATION_WORDS = ("资格", "营业执照", "信用中国", "裁判文书", "联合体", "转包", "分包", "制造商", "授权")
BINDING_WORDS = ("正本", "副本", "装订", "密封", "封装", "签字", "盖章", "电子版", "U盘", "光盘")
PRICE_WORDS = ("预算", "最高投标限价", "投标报价", "保证金", "价格分", "含税", "维保费用")


def is_full_bidding_analyst_candidate(files: list[dict[str, Any]]) -> bool:
    names = [file["filename"].lower() for file in files]
    return any("招标" in name and name.endswith(".docx") for name in names)


def pick_full_core_inputs(project_id: str, files: list[dict[str, Any]]) -> tuple[Path, Path | None] | None:
    tender: Path | None = None
    sca: Path | None = None
    for file in files:
        name = file["filename"].lower()
        path = stored_file_path(project_id, file["stored_name"])
        if "招标" in name and name.endswith(".docx") and tender is None:
            tender = path
        elif name.endswith(".docx") and "招标" not in name and sca is None:
            sca = path
    if tender:
        return tender, sca
    return None


def run_full_bidding_analyst_core(project_id: str, files: list[dict[str, Any]], report_path: Path) -> dict[str, Any] | None:
    picked = pick_full_core_inputs(project_id, files)
    if not picked:
        return None
    tender, sca = picked
    env = os.environ.copy()
    env["BIDDING_ANALYST_TENDER"] = str(tender)
    if sca:
        env["BIDDING_ANALYST_SCA"] = str(sca)
    env["BIDDING_ANALYST_OUT"] = str(report_path)
    script = Path(__file__).with_name("bidding_analyst_core.py")
    result = subprocess.run(
        [sys.executable, str(script)],
        env=env,
        cwd=str(script.parents[3]),
        capture_output=True,
        text=True,
        timeout=600,
    )
    if result.returncode != 0:
        raise RuntimeError(f"bidding-analyst 内核运行失败：{result.stderr or result.stdout}")
    try:
        summary = json.loads(result.stdout.strip().splitlines()[-1])
    except (json.JSONDecodeError, IndexError):
        summary = {
            "engine": "llm-kimi",
            "report_path": str(report_path),
            "generated_at": datetime.now().isoformat(timespec="seconds"),
        }
    summary["file_count"] = len(files)
    summary["has_solution"] = sca is not None
    risk_count = summary.pop("risk_count", 0)
    if isinstance(risk_count, int) and risk_count > 0 and not summary.get("risk_items"):
        summary["risk_items"] = [f"废标条款 {i}" for i in range(1, risk_count + 1)]
    summary.setdefault("risk_items", [])
    summary.setdefault("score_items", [])
    summary.setdefault("material_items", [])
    summary.setdefault("screenshot_items", [])
    summary.setdefault("score_simulation", {"conservative": 0, "expected": 0, "full": 0, "max": 0})
    return summary


def normalize_text(text: str) -> str:
    text = text.replace("\u3000", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_docx(path: Path) -> str:
    doc = Document(path)
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [" ".join(cell.text.split()) for cell in row.cells]
            line = " | ".join(cell for cell in cells if cell)
            if line:
                parts.append(line)
    return normalize_text("\n".join(parts))


def extract_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    return normalize_text("\n".join(page.extract_text() or "" for page in reader.pages))


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return extract_docx(path)
    if suffix == ".pdf":
        return extract_pdf(path)
    if suffix in {".txt", ".md"}:
        return normalize_text(path.read_text(encoding="utf-8", errors="ignore"))
    return ""


def esc(value: Any) -> str:
    return html.escape(str(value or ""))


def find_first(patterns: list[str], text: str) -> str:
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            value = match.group(1).strip()
            value = re.split(r"\n| {2,}|\|", value)[0].strip()
            return value[:140]
    return "待人工复核"


def lines_with(words: tuple[str, ...], text: str, limit: int, min_len: int = 8) -> list[str]:
    lines = [line.strip(" ：:;；") for line in re.split(r"[\n。；;]", text) if line.strip()]
    picked: list[str] = []
    seen = set()
    for line in lines:
        if len(line) < min_len:
            continue
        if any(word in line for word in words):
            clean = line[:320]
            if clean not in seen:
                picked.append(clean)
                seen.add(clean)
        if len(picked) >= limit:
            break
    return picked


def _classify_risk_level(text: str) -> str:
    """按四级体系分类：fatal > star > strict > review"""
    if any(w in text for w in RISK_FATAL):
        return "fatal"
    if any(w in text for w in RISK_STAR):
        return "star"
    # strict 词（应/须/必须等）必须同时命中 review 或 fatal 词才算，避免泛化误报
    has_strict = any(w in text for w in RISK_STRICT)
    has_review = any(w in text for w in RISK_REVIEW)
    has_fatal = any(w in text for w in RISK_FATAL)
    if has_strict and (has_review or has_fatal):
        return "strict"
    if has_review:
        return "review"
    return "review"


_NOISE_PATTERNS = [
    r"^\s*[目錄录]\s*[录錄次]",             # "目录" 开头
    re.compile(r"\.{4,}"),                   # 大量点号（页码填充）
    re.compile(r"^\s*[\(（]?[一二三四五六七八九十\d]+[\)）]?\s*$"),  # 纯编号行
]
_NOISE_NUM_COUNT = 5  # 一个段落里超过此数量子编号，视为目录


def _is_directory_paragraph(para: str) -> bool:
    """检测是否为目录/TOC 段落"""
    # 直接匹配目录关键词
    if re.search(r"^\s*[目錄录]\s*[录錄次]", para):
        return True
    # 大量连续点号（目录页码填充符）
    if len(re.findall(r"\.{3,}", para)) >= 3:
        return True
    # 段落内子编号过多（如 1.1.1 1.1.2 1.1.3 1.1.4 ...）
    sub_numbers = re.findall(r"\d+\.\d+(?:\.\d+)?", para)
    if len(sub_numbers) >= _NOISE_NUM_COUNT:
        return True
    # 纯编号或空白填充行
    if re.match(r"^\s*[\[\(]?\d+[\]\)\.、]?\s*(\.{2,})?\s*$", para) and len(para) < 40:
        return True
    return False


RISK_LABEL = {"fatal": "致命", "star": "星号否决", "strict": "严重", "review": "需注意"}


def extract_risk_detail(combined: str, limit: int = 20) -> list[dict[str, str]]:
    """参照 bidding-analyst skill 提取废标条款，返回结构化数据。
    比原 lines_with 升级点：
    1. 按段落拆分而非按行
    2. 四级风险分类（致命/星号/严重/需注意）
    3. 过滤目录/TOC 噪声段落
    4. 去重
    """
    paras = re.split(r"\n{2,}", combined)
    items: list[dict[str, str]] = []
    seen = set()

    for para in paras:
        clean = para.strip()
        if len(clean) < 15:
            continue
        # 过滤目录/TOC 噪声
        if _is_directory_paragraph(clean):
            continue
        # 快速检查是否有任何风险关键词
        if not any(w in clean for w in KEY_RISK_WORDS):
            continue
        level = _classify_risk_level(clean)
        # 去重：长文本取前 200 字做特征
        sig = clean[:200]
        if sig in seen:
            continue
        seen.add(sig)
        # 智能截取：长段落取前 320 字并用换行分隔，短段落全取
        if len(clean) <= 320:
            context = clean
        else:
            # 在 300-340 字之间找最近的句号或换行处截断
            cut = clean[:340].rfind("。")
            if cut > 200:
                context = clean[:cut + 1] + "…"
            else:
                context = clean[:320] + "…"
        items.append({"text": context, "level": level})
        if len(items) >= limit:
            break

    # 按风险等级排序：star > fatal > strict > review
    rank = {"star": 0, "fatal": 1, "strict": 2, "review": 3}
    items.sort(key=lambda x: rank.get(x["level"], 9))
    return items


def detect_project_name(text: str, fallback: str) -> str:
    value = find_first([r"项目名称[:：]\s*([^\n]+)", r"招标项目名称[:：]\s*([^\n]+)"], text)
    return fallback if value == "待人工复核" else value


def detect_dates(text: str) -> list[dict[str, str]]:
    raw_items = [
        ("招标文件领取", r"领取时间[:：]?\s*([0-9]{4}年[0-9]{1,2}月[0-9]{1,2}日[^。\n]*)"),
        ("投标截止/开标", r"投标.*?截止.*?开标时间[:：]?\s*([0-9]{4}年[0-9]{1,2}月[0-9]{1,2}日[0-9]{1,2}时[0-9]{1,2}分[^。\n]*)"),
        ("开标时间", r"开标时间[:：]?\s*([0-9]{4}年[0-9]{1,2}月[0-9]{1,2}日[0-9]{1,2}时[0-9]{1,2}分[^。\n]*)"),
        ("投标有效期", r"投标有效期[:：]?\s*([^\n。；]{4,80})"),
        ("保证金", r"保证金[^\n。；]{0,40}([0-9]{4}年[0-9]{1,2}月[0-9]{1,2}日[^\n。；]*)"),
    ]
    items: list[dict[str, str]] = []
    for name, pattern in raw_items:
        match = re.search(pattern, text)
        if not match:
            continue
        date_text = match.group(1).strip()
        status = "需人工确认"
        date_match = re.search(r"([0-9]{4})年([0-9]{1,2})月([0-9]{1,2})日", date_text)
        if date_match:
            y, m, d = map(int, date_match.groups())
            dt = date(y, m, d)
            status = f"已截止（截至 {TODAY.isoformat()} 已过去 {(TODAY - dt).days} 天）" if dt < TODAY else f"剩余 {(dt - TODAY).days} 天"
        items.append({"name": name, "date": date_text, "status": status, "note": "请以招标文件原文和平台公告为准"})
    if not items:
        items.append({"name": "关键时间节点", "date": "待人工复核", "status": "未自动识别", "note": "建议人工核对公告、投标人须知和平台通知"})
    return items[:8]


def build_summary(project_id: str, files: list[dict[str, Any]], extracted: dict[str, str]) -> dict[str, Any]:
    combined = "\n".join(extracted.values())
    fallback = files[0]["filename"].rsplit(".", 1)[0] if files else f"项目 {project_id}"
    risk_detail = extract_risk_detail(combined, 20)
    risk_items = [item["text"] for item in risk_detail]  # 兼容前端 length 计数
    score_items = lines_with(SCORE_WORDS, combined, 18)
    material_items = lines_with(MATERIAL_WORDS, combined, 24)
    screenshot_items = lines_with(SCREENSHOT_WORDS, combined, 18)
    qualification_items = lines_with(QUALIFICATION_WORDS, combined, 18)
    price_items = lines_with(PRICE_WORDS, combined, 16)
    binding_items = lines_with(BINDING_WORDS, combined, 16)
    return {
        "project_name": detect_project_name(combined, fallback),
        "project_no": find_first([r"项目编号[:：]\s*([^\n]+)", r"招标编号[:：]\s*([^\n]+)"], combined),
        "buyer": find_first([r"招\s*标\s*人[:：]\s*([^\n]+)", r"采购人[:：]\s*([^\n]+)"], combined),
        "agent": find_first([r"招标代理[:：]\s*([^\n]+)", r"代理机构[:：]\s*([^\n]+)"], combined),
        "budget": find_first([r"项目预算[:：]\s*([^\n]+)", r"最高投标限价[:：]\s*([^\n]+)", r"预算金额[:：]\s*([^\n]+)"], combined),
        "deadline": find_first([r"投标.*?截止.*?[:：]\s*([^\n]+)", r"开标时间[:：]\s*([^\n]+)"], combined),
        "file_count": len(files),
        "text_stats": dict(Counter({name: len(text) for name, text in extracted.items()})),
        "risk_items": risk_items,
        "score_items": score_items,
        "material_items": material_items,
        "screenshot_items": screenshot_items,
        "qualification_items": qualification_items,
        "price_items": price_items,
        "binding_items": binding_items,
        "timeline_items": detect_dates(combined),
        "has_solution": len(files) > 1,
        "risk_detail": risk_detail,  # 新增：四级分类风险数据
        "score_simulation": {"conservative": 0, "expected": 0, "full": 0, "max": 0},
        "generated_at": datetime.now().isoformat(timespec="seconds"),
    }


def list_cards(items: list[str], empty: str, cls: str = "info-card", checkbox: bool = False) -> str:
    if not items:
        return f'<p class="muted">{esc(empty)}</p>'
    cards = []
    for index, item in enumerate(items, 1):
        check = '<label class="check"><input type="checkbox"> 已准备</label>' if checkbox else ""
        cards.append(
            f'<article class="{cls}"><div class="card-head"><span class="num">{index}</span>{check}</div>'
            f'<p>{esc(item)}</p></article>'
        )
    return '<div class="card-list">' + "\n".join(cards) + "</div>"


def render_risk_cards(risk_detail: list[dict[str, str]], empty: str = "暂未自动识别到明显废标条款，请人工复核投标人须知和评标办法。") -> str:
    """四级风险彩色卡片渲染。🟥 star > 🟥 fatal > 🟧 strict > 🟨 review"""
    if not risk_detail:
        return f'<div class="info-block"><p class="muted">{esc(empty)}</p></div>'
    cls_map = {"star": "risk-star", "fatal": "risk-fatal", "strict": "risk-strict", "review": "risk-review"}
    guidance = "以下条目由关键词匹配自动提取，未经大模型语义分析。请以招标文件原文为准，逐条人工核验。"
    cards = []
    for index, item in enumerate(risk_detail, 1):
        level = item.get("level", "review")
        cls = cls_map.get(level, "risk-review")
        label = RISK_LABEL.get(level, "需注意")
        tag_cls = f"level-{level}" if level in ("star", "fatal", "strict", "review") else "level-review"
        prefix = "⚠️ 星号否决：" if level == "star" else ("🔴 致命废标：" if level == "fatal" else "")
        text = item["text"]
        # 长文本用 details/summary 折叠，短文本直接展示
        if len(text) > 180:
            summary_text = text[:160].rstrip("。；;，,、 ") + "…"
            body = (
                f'<article class="{cls}"><div class="card-head">'
                f'<span class="num">{index}</span>'
                f'<span class="risk-level-tag {tag_cls}">{label}</span>'
                f'</div><details><summary>{prefix}{esc(summary_text)}</summary>'
                f'<div class="risk-full-text">{esc(text)}</div></details></article>'
            )
        else:
            body = (
                f'<article class="{cls}"><div class="card-head">'
                f'<span class="num">{index}</span>'
                f'<span class="risk-level-tag {tag_cls}">{label}</span>'
                f'</div><p>{prefix}{esc(text)}</p></article>'
            )
        cards.append(body)
    return (
        f'<p class="section-hint">{guidance}</p>'
        + '<div class="card-list">' + "\n".join(cards) + "</div>"
    )


def timeline(items: list[dict[str, str]]) -> str:
    rendered = []
    for item in items:
        rendered.append(
            f'<article class="timeline-item"><span></span><div><strong>{esc(item["name"])}</strong>'
            f'<p>{esc(item["date"])}</p><em>{esc(item["status"])}</em><small>{esc(item["note"])}</small></div></article>'
        )
    return '<div class="timeline">' + "\n".join(rendered) + "</div>"


def score_bar(label: str, value: int, total: int, note: str) -> str:
    width = min(100, round(value / total * 100)) if total else 0
    return (
        f'<div class="score-row"><span>{esc(label)}</span><div class="bar"><i style="width:{width}%"></i></div>'
        f'<strong>{value}/{total} 分</strong><small>{esc(note)}</small></div>'
    )


def simulation_cards(summary: dict[str, Any]) -> str:
    sim = summary["score_simulation"]
    if not summary["has_solution"]:
        return '<p class="muted">未上传公司方案参数，暂不进行我方方案得分测算。</p>'
    items = [
        ("技术保守得分", sim["conservative"], "只计入已能从方案文件直接支撑的能力。"),
        ("技术预期得分", sim["expected"], "假设投标前补齐截图、承诺和功能证明材料。"),
        ("技术满分潜力", sim["full"], "需要完整响应评分表全部技术条款并提供可核验证据。"),
    ]
    blocks = [
        f'<div class="metric"><span>{esc(name)}</span><strong>{score}/{sim["max"]}</strong><p>{esc(note)}</p></div>'
        for name, score, note in items
    ]
    detail_items = [
        "语言、包管理器、二进制格式、镜像、SBOM、组件库规模等评分项需逐项截图证明。",
        "当前测算按“缺材料不默认满分”原则处理；未见明确证明的评分项保守计 0 或待确认。",
        "价格分未计算：需要最终报价及评标基准价或其他投标报价场景。",
    ]
    return '<div class="metric-grid">' + "".join(blocks) + "</div>" + list_cards(detail_items, "", "score-card")


def framework_html() -> str:
    sections = [
        ("一、商务响应文件", ["1.1 投标函", "1.2 法定代表人身份证明", "1.3 法定代表人授权委托书", "1.4 投标保证金证明", "1.5 资格证明文件"]),
        ("二、报价文件", ["2.1 开标一览表", "2.2 分项报价表", "2.3 维保费用说明"]),
        ("三、技术响应文件", ["3.1 技术规格响应表", "3.2 平台功能说明", "3.3 部署实施方案", "3.4 测试截图与功能证明", "3.5 售后服务方案"]),
        ("四、评分索引与证明材料", ["4.1 评分索引表", "4.2 业绩证明材料", "4.3 认证证书与资质材料", "4.4 其他补充资料"]),
    ]
    html_parts = []
    for title, subs in sections:
        html_parts.append(f"<h2>{esc(title)}</h2>")
        html_parts.extend(f"<h3>{esc(sub)}</h3>" for sub in subs)
    return "\n".join(html_parts)


def checklist_section(risks: list[str]) -> str:
    risk_checks = risks[:8] or ["所有废标条款已人工核验"]
    groups = {
        "A. 废标条款核查": [f"{item[:80]} 已满足" for item in risk_checks],
        "B. 内容完整性": ["商务文件、技术文件、报价文件均已齐套", "投标文件目录与正文页码一致"],
        "C. 数据一致性": ["项目名称、编号、报价、交期、质保期前后一致", "电子版与纸质版内容一致"],
        "D. 签章完整性": ["投标函、授权书、报价表、偏离表已按要求签字盖章", "授权代表签署材料已附授权文件"],
        "E. 截图与证明材料": ["功能截图、查询截图、证书扫描件、合同关键页已归档", "评分索引能定位到对应证明材料"],
    }
    blocks = []
    for title, items in groups.items():
        labels = "".join(f'<label class="check"><input type="checkbox"> {esc(item)}</label>' for item in items)
        blocks.append(f"<details open><summary>{esc(title)}</summary>{labels}</details>")
    return '<div class="progress"><i id="check-progress"></i><span id="check-count">0/0</span></div>' + "".join(blocks)


def render_report(summary: dict[str, Any], files: list[dict[str, Any]]) -> str:
    css = """
    :root{--bg:#f5f7fb;--panel:#fff;--text:#172033;--muted:#667085;--line:#d9e0ea;--primary:#0f766e;--primary-2:#115e59;--accent:#c2410c;--danger:#b42318;--warn:#b54708;--ok:#067647;--soft:#eef6f5}
    *{box-sizing:border-box}html{scroll-behavior:smooth}body{margin:0;background:var(--bg);color:var(--text);font-family:-apple-system,BlinkMacSystemFont,"Segoe UI",sans-serif;letter-spacing:0}a{color:inherit}
    .report-shell{display:grid;grid-template-columns:260px minmax(0,1fr);min-height:100vh}.toc{position:sticky;top:0;height:100vh;padding:22px;background:#0b1220;color:#d7dee9;overflow:auto}.toc h1{font-size:18px;margin:0 0 8px;color:#fff}.toc p{font-size:13px;color:#98a2b3;line-height:1.55}.toc a{display:block;padding:9px 10px;margin:4px 0;border-radius:6px;text-decoration:none;font-size:13px}.toc a:hover{background:rgba(255,255,255,.08)}
    main{max-width:1220px;width:100%;padding:28px;margin:0 auto}.hero{padding:26px;border:1px solid var(--line);border-radius:8px;background:linear-gradient(135deg,#fff,#f2fbf9)}.hero h2{margin:0 0 10px;font-size:28px}.hero p{margin:0;color:var(--muted)}
    section{margin:18px 0;padding:22px;border:1px solid var(--line);border-radius:8px;background:var(--panel)}.section-title{display:flex;justify-content:space-between;gap:16px;align-items:center;margin-bottom:16px}h2{margin:0;font-size:20px}h3{font-size:16px;margin:0 0 10px}p,li{line-height:1.65}.muted{color:var(--muted)}
    .metric-grid,.overview{display:grid;grid-template-columns:repeat(4,minmax(0,1fr));gap:12px}.metric,.overview div{padding:14px;border:1px solid var(--line);border-radius:8px;background:#fbfcfd}.metric span,.overview span{display:block;color:var(--muted);font-size:13px}.metric strong{display:block;margin-top:8px;font-size:24px}.metric p{margin:8px 0 0;color:var(--muted);font-size:13px}
    .badge{display:inline-flex;align-items:center;border-radius:999px;padding:4px 9px;font-size:12px;font-weight:700}.fatal{background:#fee4e2;color:var(--danger)}.ok{background:#dcfae6;color:var(--ok)}.warn{background:#fef0c7;color:var(--warn)}
    .timeline{position:relative}.timeline-item{display:grid;grid-template-columns:22px minmax(0,1fr);gap:10px;padding:10px 0}.timeline-item>span{width:12px;height:12px;border-radius:999px;background:var(--primary);margin-top:7px}.timeline-item div{border-bottom:1px solid var(--line);padding-bottom:10px}.timeline-item p{margin:4px 0}.timeline-item em{display:block;color:var(--warn);font-style:normal;font-weight:700}.timeline-item small{color:var(--muted)}
    .card-list{display:grid;grid-template-columns:repeat(2,minmax(0,1fr));gap:12px}.info-card,.risk-card,.risk-star,.risk-fatal,.risk-strict,.risk-review,.score-card{border:1px solid var(--line);border-radius:8px;background:#fbfcfd;padding:14px}.risk-card{border-left:4px solid var(--danger)}.risk-star{border:2px solid #b42318;border-left:5px solid #b42318;background:#fff5f5}.risk-fatal{border-left:4px solid #b42318;background:#fef2f2}.risk-strict{border-left:4px solid #b54708;background:#fff8f1}.risk-review{border-left:3px solid #92400e;background:#fffbf5}.card-head{display:flex;justify-content:space-between;align-items:center;gap:10px}.num{display:inline-grid;place-items:center;width:28px;height:28px;border-radius:999px;background:var(--soft);color:var(--primary);font-weight:800}.risk-star .num{background:#fecaca;color:#b42318}.risk-fatal .num{background:#fee2e2;color:#b42318}.risk-strict .num{background:#fed7aa;color:#b54708}.risk-review .num{background:#fef3c7;color:#92400e}.risk-level-tag{display:inline-flex;align-items:center;border-radius:999px;padding:2px 8px;font-size:11px;font-weight:700;margin-left:8px}.level-star{background:#fecaca;color:#b42318}.level-fatal{background:#fee2e2;color:#b42318}.level-strict{background:#fed7aa;color:#b54708}.level-review{background:#fef3c7;color:#92400e}
    .risk-full-text{padding:8px 0 0;font-size:14px;color:#475569;line-height:1.7}
    .info-block{padding:10px 14px;border-radius:8px;background:#f0f9ff;border:1px solid #bae6fd;margin-bottom:14px}
    .section-hint{font-size:13px;color:#475569;margin:0 0 8px;line-height:1.6}
    .scoring-grid{display:grid;grid-template-columns:minmax(0,1fr) minmax(0,1.6fr);gap:18px;align-items:start}
    .scoring-col{border:1px solid var(--line);border-radius:8px;padding:16px;background:#fbfcfd}
    .scoring-col h3{margin-top:0}
    .scoring-main{margin-bottom:14px}
    .scoring-main h3{margin:0 0 10px;font-size:15px}
    .ref-block{border:1px solid #d1d5db;border-radius:8px;padding:10px 14px;background:#fafbfc;margin-top:10px}
    .ref-block summary{cursor:pointer;font-weight:700;font-size:14px;color:#475569}
    .ref-block table{margin-top:8px}
    .scoring-table th,.scoring-table td,.check-table th,.check-table td{border-bottom:1px solid #e2e8f0;padding:8px 10px;text-align:left;font-size:13px}
    .scoring-table th,.check-table th{background:#f8fafc;color:var(--muted);font-weight:700}
    .risk-full-text + details summary{cursor:pointer;font-size:14px;line-height:1.6}
    .risk-star details summary,.risk-fatal details summary,.risk-strict details summary,.risk-review details summary{font-weight:700}
    details .risk-full-text{margin-top:8px;padding:10px 12px;background:rgba(255,255,255,.6);border-radius:6px;border:1px solid var(--line)}
    .score-row{display:grid;grid-template-columns:120px minmax(0,1fr)90px;gap:12px;align-items:center;margin:12px 0}.score-row small{grid-column:2/4;color:var(--muted)}.bar{height:10px;border-radius:999px;background:#e7edf3;overflow:hidden}.bar i{display:block;height:100%;background:linear-gradient(90deg,var(--primary),#14b8a6)}
    details{border:1px solid var(--line);border-radius:8px;padding:12px;background:#fbfcfd;margin:10px 0}summary{cursor:pointer;font-weight:800}.check{display:flex;gap:8px;align-items:flex-start;margin:10px 0;color:#344054}.check input{margin-top:4px}.progress{display:flex;align-items:center;gap:10px;background:#edf2f7;border-radius:999px;height:12px;margin-bottom:14px}.progress i{display:block;height:100%;width:0;border-radius:999px;background:var(--primary)}.progress span{font-size:12px;color:var(--muted)}
    table{width:100%;border-collapse:collapse;table-layout:fixed}th,td{border-bottom:1px solid var(--line);padding:10px;text-align:left;vertical-align:top;word-break:break-word}th{color:var(--muted);background:#f8fafc}.word-framework{padding:14px;border:1px dashed var(--line);border-radius:8px;background:#fbfcfd}.word-framework h2{font-size:18px;margin:14px 0 8px}.word-framework h3{font-size:15px;margin:8px 0}
    button,.button{min-height:38px;border:0;border-radius:6px;padding:0 14px;background:var(--primary);color:#fff;font-weight:800;text-decoration:none;cursor:pointer}button:hover,.button:hover{background:var(--primary-2)}
    @media(max-width:980px){.report-shell{grid-template-columns:1fr}.toc{position:relative;height:auto}.toc nav{display:grid;grid-template-columns:repeat(2,minmax(0,1fr));gap:4px}main{padding:14px}.metric-grid,.overview,.card-list,.scoring-grid{grid-template-columns:1fr}.score-row{grid-template-columns:1fr}.score-row small{grid-column:auto}}
    @media print{.toc,button{display:none}.report-shell{display:block}main{max-width:none;padding:0}section,.hero{break-inside:avoid;background:#fff}}
    """
    js = """
    function updateChecks(){const checks=[...document.querySelectorAll('input[type="checkbox"]')];const done=checks.filter(x=>x.checked).length;const total=checks.length;const pct=total?Math.round(done/total*100):0;const bar=document.getElementById('check-progress');const count=document.getElementById('check-count');if(bar)bar.style.width=pct+'%';if(count)count.textContent=done+'/'+total}
    document.addEventListener('change',updateChecks);updateChecks();
    document.querySelectorAll('[data-copy-target]').forEach(btn=>btn.addEventListener('click',async()=>{const el=document.getElementById(btn.dataset.copyTarget);const text=el.innerText;try{if(navigator.clipboard&&window.isSecureContext){await navigator.clipboard.writeText(text)}else{const area=document.createElement('textarea');area.value=text;area.style.position='fixed';area.style.opacity='0';document.body.appendChild(area);area.select();document.execCommand('copy');area.remove()}btn.textContent='已复制'}catch(e){btn.textContent='请手动复制'}}));
    """
    chapters = [
        ("overview", "一、项目基本信息速览"),
        ("timeline", "二、关键时间节点"),
        ("risks", "三、废标条款清单"),
        ("scoring", "四、评分标准详解"),
        ("score-simulation", "五、我方方案得分测算"),
        ("qualification", "六、资质门槛清单"),
        ("price", "七、价格限制与报价策略参考"),
        ("bonus", "八、加分机会与满分材料清单"),
        ("evidence-materials", "九、截图与证明材料清单"),
        ("binding", "十、装订封装要求"),
        ("bid-framework", "十一、投标文件框架"),
        ("checklist", "十二、投标文件自检 Checklist"),
        ("todo", "十三、素材准备 TodoList"),
        ("score-actions", "十四、得分补强行动清单"),
        ("advice", "十五、综合分析与建议"),
    ]
    nav = "".join(f'<a href="#{cid}">{esc(title)}</a>' for cid, title in chapters)
    file_rows = "".join(f"<tr><td>{esc(f['filename'])}</td><td>{f['size']}</td><td>{esc(f['uploaded_at'])}</td></tr>" for f in files)
    return f"""<!doctype html><html lang="zh-CN"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width, initial-scale=1"><title>{esc(summary["project_name"])} 招标文件分析报告</title><style>{css}</style></head>
<body><div class="report-shell"><aside class="toc"><h1>标书分析报告</h1><p>{esc(summary["project_name"])}<br>生成时间：{esc(summary["generated_at"])}</p><nav>{nav}</nav></aside><main>
<header class="hero"><h2>{esc(summary["project_name"])} 招标文件分析报告</h2><p>分析人：标书阅读专员 AI · 报告包含废标、评分、得分测算、材料清单、装订要求、框架和自检工具。</p></header>
<section id="overview"><div class="section-title"><h2>一、项目基本信息速览</h2><span class="badge ok">已解析</span></div><div class="overview"><div><span>项目编号</span>{esc(summary["project_no"])}</div><div><span>招标/采购人</span>{esc(summary["buyer"])}</div><div><span>招标代理</span>{esc(summary["agent"])}</div><div><span>预算/限价</span>{esc(summary["budget"])}</div></div><h3>上传文件</h3><table><thead><tr><th>文件名</th><th>大小 Byte</th><th>上传时间</th></tr></thead><tbody>{file_rows}</tbody></table></section>
<section id="timeline"><h2>二、关键时间节点</h2>{timeline(summary["timeline_items"])}</section>
<section id="risks"><div class="section-title"><h2>三、废标条款清单</h2><span class="badge fatal">⚠️ 优先核验</span></div>{render_risk_cards(summary.get("risk_detail", []))}</section>
<section id="scoring"><div class="section-title"><h2>四、评分标准详解</h2><span class="badge warn">参照招标文件原文</span></div><div class="info-block"><p class="section-hint">🎯 <strong>怎么读：</strong>下方为从招标文件中自动提取的评分规则线索（主要信息）。底部为常见评分大类结构参考，请以招标文件「评标办法」「评分标准」章节为准。</p></div><div class="scoring-main"><h3>📋 提取到的评分规则线索</h3>{list_cards(summary["score_items"], "暂未识别到评分标准线索，请人工查阅招标文件「评标办法」「评分标准」章节。", "score-card")}</div><details class="ref-block"><summary>📊 常见评分大类结构参考（以招标文件原文为准）</summary><table class="scoring-table"><thead><tr><th>评分大类</th><th>参考权重</th><th>参考满分</th></tr></thead><tbody><tr><td>技术部分</td><td>50-60%</td><td>50-60 分</td></tr><tr><td>商务部分</td><td>10-20%</td><td>10-20 分</td></tr><tr><td>价格部分</td><td>20-40%</td><td>20-40 分</td></tr></tbody></table></details></section>
<section id="score-simulation"><h2>五、我方方案得分测算</h2>{simulation_cards(summary)}</section>
<section id="qualification"><div class="section-title"><h2>六、资质门槛清单</h2><span class="badge warn">硬性门槛</span></div><div class="info-block"><p class="section-hint">📋 <strong>怎么读：</strong>下方为从招标文件中自动提取的资质/资格相关线索（主要信息）。每条请对照招标文件「投标人资格要求」「资格审查」章节确认。底部为人工核查清单参考。</p></div><div class="scoring-main"><h3>📌 已识别的资质线索</h3>{list_cards(summary["qualification_items"], "暂未识别到资质门槛线索，请人工查阅招标文件「投标人资格要求」。", "info-card", True)}</div><details class="ref-block"><summary>🔍 人工核查清单参考（请以招标文件原文为准）</summary><table class="check-table"><thead><tr><th>核查项</th><th>状态</th><th>要求</th></tr></thead><tbody><tr><td>营业执照经营范围</td><td><span class="badge warn">待确认</span></td><td>是否覆盖本项目采购内容</td></tr><tr><td>行业资质证书</td><td><span class="badge warn">待确认</span></td><td>核对等级要求（如：甲级/乙级</td></tr><tr><td>类似项目业绩</td><td><span class="badge warn">待确认</span></td><td>数量、金额门槛、时间范围</td></tr><tr><td>项目经理资质</td><td><span class="badge warn">待确认</span></td><td>证书类型、级别、社保要求</td></tr><tr><td>信用要求</td><td><span class="badge warn">待确认</span></td><td>信用中国截图、无违法记录</td></tr><tr><td>财务要求</td><td><span class="badge warn">待确认</span></td><td>审计报告年限、资产负债率</td></tr></tbody></table></details></section>
<section id="price"><div class="section-title"><h2>七、价格限制与报价策略参考</h2><span class="badge warn">报价策略</span></div><div class="info-block"><p class="section-hint">💰 <strong>怎么读：</strong>下方为从招标文件中提取的价格相关信息（预算/限价/保证金/报价方式等，主要信息）。底部为常见价格分计算公式参考（以招标文件原文为准）。</p></div><div class="scoring-main"><h3>📌 价格相关信息</h3>{list_cards(summary["price_items"], "暂未识别到价格限制线索，请人工查阅「投标人须知」和「报价要求」。", "info-card")}</div><details class="ref-block"><summary>📐 常见价格分计算公式参考（以招标文件原文为准）</summary><table class="scoring-table"><thead><tr><th>计算公式</th><th>场景</th></tr></thead><tbody><tr><td>低价优先法</td><td>最低价满分，其余按比例递减</td></tr><tr><td>均价基准法</td><td>接近平均价得高分</td></tr><tr><td>综合评分法</td><td>多因素加权</td></tr></tbody></table></details><blockquote>⚠️ 报价策略建议属于分析推断。最终报价需结合成本、竞争态势和价格分计算公式人工决策。</blockquote></section>
<section id="bonus"><h2>八、加分机会与满分材料清单</h2>{list_cards(summary["score_items"][:10], "暂未识别到明确加分机会。", "info-card", True)}</section>
<section id="evidence-materials"><h2>九、截图与证明材料清单</h2><details open><summary>9.1 需提供截图的条款</summary>{list_cards(summary["screenshot_items"], "暂未识别到明确截图条款。", "info-card", True)}</details><details open><summary>9.2 非截图类证明材料</summary>{list_cards(summary["material_items"], "暂未识别到明确证明材料条款。", "info-card", True)}</details></section>
<section id="binding"><div class="section-title"><h2>十、装订封装要求</h2><span class="badge warn">格式规范</span></div><div class="info-block"><p class="section-hint">📦 <strong>怎么读：</strong>下方为从招标文件中提取的装订封装线索（主要信息）。底部为提交前必查清单参考，请以招标文件原文为准。</p></div><div class="scoring-main"><h3>📌 提取到的装订封装线索</h3>{list_cards(summary["binding_items"], "暂未识别到装订封装要求，请人工查阅「投标人须知」中关于投标文件装订、密封、递交的章节。", "info-card")}</div><details class="ref-block"><summary>📋 提交前必查清单参考（请以招标文件原文为准）</summary><table class="check-table"><thead><tr><th>核查项</th><th>状态</th></tr></thead><tbody><tr><td>正本份数</td><td><span class="badge warn">待确认</span></td></tr><tr><td>副本份数</td><td><span class="badge warn">待确认</span></td></tr><tr><td>装订方式（胶装/线装）</td><td><span class="badge warn">待确认</span></td></tr><tr><td>密封要求（封条/密封章）</td><td><span class="badge warn">待确认</span></td></tr><tr><td>电子版介质（U盘/光盘）</td><td><span class="badge warn">待确认</span></td></tr><tr><td>封套标识信息</td><td><span class="badge warn">待确认</span></td></tr><tr><td>签字盖章位置</td><td><span class="badge warn">待确认</span></td></tr></tbody></table></details></section>
<section id="bid-framework"><div class="section-title"><h2>十一、投标文件框架</h2><button type="button" data-copy-target="framework-content">复制框架</button></div><div id="framework-content" class="word-framework">{framework_html()}</div></section>
<section id="checklist"><h2>十二、投标文件自检 Checklist</h2>{checklist_section(summary["risk_items"])}</section>
<section id="todo"><h2>十三、素材准备 TodoList</h2>{list_cards(summary["material_items"][:18], "暂无素材任务。", "info-card", True)}</section>
<section id="score-actions"><h2>十四、得分补强行动清单</h2><div class="card-list"><article class="info-card"><div class="card-head"><span class="badge fatal">立即处理</span><label class="check"><input type="checkbox"> 完成</label></div><h3>合规与星号条款复核</h3><p>关联评分/风险项：废标条款、资格审查、重要技术条款。补强动作：逐条核验响应表、保证金、授权、签字盖章和报价上限。</p></article><article class="info-card"><div class="card-head"><span class="badge warn">高收益补强</span><label class="check"><input type="checkbox"> 完成</label></div><h3>功能截图与评分证据</h3><p>关联评分项：技术功能、检测能力、SBOM、镜像、组件库、包管理器。补强动作：按评分项建立截图索引，补齐可核验页面截图。</p></article><article class="info-card"><div class="card-head"><span class="badge warn">中收益补强</span><label class="check"><input type="checkbox"> 完成</label></div><h3>方案章节完整度</h3><p>关联评分项：平台部署、运维、后期运营、应急响应。补强动作：补充场景化说明、实施计划、服务承诺和应急流程。</p></article><article class="info-card"><div class="card-head"><span class="badge ok">可选择</span><label class="check"><input type="checkbox"> 完成</label></div><h3>低确定性分值</h3><p>关联评分项：依赖主观评价或外部报价的分值。处理建议：标注假设，交由人工结合竞争态势决策。</p></article></div></section>
<section id="advice"><h2>十五、综合分析与建议</h2><div class="card-list"><article class="info-card"><h3>先保合规</h3><p>优先核验废标条款、资格门槛、星号条款、保证金、签字盖章和报价上限。</p></article><article class="info-card"><h3>再补得分</h3><p>围绕评分项逐条补齐截图、证书、承诺、功能说明和评分索引，避免“有能力但无证明”。</p></article></div><blockquote>本报告由 AI 辅助生成，所有关键条款请以招标文件原文为准，建议人工复核后使用。</blockquote></section>
</main></div><script>{js}</script></body></html>"""


def archive_report(project_id: str, report_path: Path, analyze_type: str, engine: str | None) -> str:
    """把当前报告按时间戳归档，返回归档后的文件名。"""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    archive_name = f"report_{timestamp}.html"
    archive_path = report_path.parent / archive_name
    if report_path.exists():
        shutil.copy2(report_path, archive_path)
    add_report(
        project_id=project_id,
        filename=archive_name,
        stored_name=archive_name,
        analyze_type=analyze_type,
        engine=engine,
        size=archive_path.stat().st_size if archive_path.exists() else 0,
    )
    return archive_name


def analyze_project(project_id: str, analyze_type: str = "general") -> None:
    project_dir = PROJECTS_DIR / project_id
    extracted_dir = project_dir / "extracted"
    reports_dir = project_dir / "reports"
    extracted_dir.mkdir(parents=True, exist_ok=True)
    reports_dir.mkdir(parents=True, exist_ok=True)
    update_project(project_id, status="running", error=None)
    try:
        files = list_files(project_id)
        report_path = reports_dir / "analysis-report.html"
        if analyze_type == "advanced" and is_full_bidding_analyst_candidate(files):
            summary = run_full_bidding_analyst_core(project_id, files, report_path)
            if summary:
                archive_report(project_id, report_path, "advanced", summary.get("engine"))
                zip_path = reports_dir / "result-package.zip"
                with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
                    zf.write(report_path, "analysis-report.html")
                    for upload_file in (project_dir / "uploads").glob("*"):
                        if upload_file.is_file():
                            zf.write(upload_file, f"uploads/{upload_file.name}")
                update_project(project_id, status="done", summary=summary, report_path=str(report_path), error=None)
                return

        extracted: dict[str, str] = {}
        for file in files:
            path = stored_file_path(project_id, file["stored_name"])
            text = extract_text(path)
            if text:
                extracted[file["filename"]] = text
                (extracted_dir / f"{file['stored_name']}.txt").write_text(text, encoding="utf-8")
        if not extracted:
            raise ValueError("未能从上传文件中解析出文本，请确认文件格式为 docx、pdf、txt 或 md。")
        summary = build_summary(project_id, files, extracted)
        report_path.write_text(render_report(summary, files), encoding="utf-8")
        archive_report(project_id, report_path, "general", "keyword-extract")
        zip_path = reports_dir / "result-package.zip"
        with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
            zf.write(report_path, "analysis-report.html")
            for text_file in extracted_dir.glob("*.txt"):
                zf.write(text_file, f"extracted/{text_file.name}")
        update_project(project_id, status="done", summary=summary, report_path=str(report_path), error=None)
    except Exception as exc:
        log_path = project_dir / "logs"
        log_path.mkdir(parents=True, exist_ok=True)
        (log_path / "latest-error.txt").write_text(str(exc), encoding="utf-8")
        update_project(project_id, status="failed", error=str(exc))
