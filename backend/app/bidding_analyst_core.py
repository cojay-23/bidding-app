#!/usr/bin/env python3
"""bidding_analyst_core.py — 调用大模型分析招标文件，输出 HTML 报告。

被 analyzer.py 通过 subprocess 拉起，接收环境变量：
  BIDDING_ANALYST_TENDER  必需，招标文件路径（docx/pdf/txt/md）
  BIDDING_ANALYST_SCA     可选，方案/产品材料路径
  BIDDING_ANALYST_OUT     必需，HTML 报告输出路径
  MOONSHOT_API_KEY        必需，Kimi/Moonshot API 密钥
  LLM_BASE_URL            可选，默认 https://api.moonshot.cn/v1
  LLM_MODEL               可选，默认 kimi-k2.6（招标文件分析推荐 kimi-k2.6；kimi-k2.7-code 为编程专用模型）

成功时：HTML 写到 OUT，stdout 输出一行 JSON 摘要供 analyzer.py 解析。
失败时：错误信息写到 stderr，非零退出码。
"""
from __future__ import annotations

import json
import os
import re
import sys
from pathlib import Path

import bleach
from bleach.css_sanitizer import CSSSanitizer
from docx import Document
from openai import BadRequestError, OpenAI
from pypdf import PdfReader


API_KEY = (os.environ.get("MOONSHOT_API_KEY") or os.environ.get("DEEPSEEK_API_KEY") or "").strip()
BASE_URL = os.environ.get("LLM_BASE_URL", "https://api.moonshot.cn/v1").strip()
MODEL = os.environ.get("LLM_MODEL", "kimi-k2.6").strip()

MAX_CHARS = 120000


def extract_docx(path: Path) -> str:
    doc = Document(path)
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [" ".join(cell.text.split()) for cell in row.cells]
            line = " | ".join(cell for cell in cells if cell)
            if line:
                parts.append(line)
    return "\n".join(parts)


def extract_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return extract_docx(path)
    if suffix == ".pdf":
        return extract_pdf(path)
    if suffix in {".txt", ".md"}:
        return path.read_text(encoding="utf-8", errors="ignore")
    return ""


def truncate(text: str, limit: int = MAX_CHARS) -> str:
    if len(text) <= limit:
        return text
    return text[:limit] + "\n\n[注：文件内容过长，已截断，仅分析前部分内容]"


def build_prompt(tender_text: str, solution_text: str | None) -> str:
    solution_block = ""
    if solution_text:
        solution_block = f"\n\n## 方案/产品材料（供参考，用于评估我方匹配度）\n\n{truncate(solution_text, 60000)}\n"

    return f"""你是一位拥有 10 年招投标经验的资深标书分析师。你的核心能力是从几十到几百页的招标文件中精准提取关键信息、识别所有可能导致废标的陷阱条款、将评分标准拆解为可执行的得分策略、输出立即可用的分析报告。

必须完整阅读整个招标文件，不要只看前几页。评分标准、废标条款往往散落在文件的各个章节。

## 招标文件内容

{truncate(tender_text)}
{solution_block}

## 结构化信息提取（八大模块）

按以下八大模块逐项提取信息，宁可多提不可遗漏。报告组织顺序必须遵循“基础信息靠前、细节信息靠后”的阅读逻辑。

### 模块一：项目基本信息

提取项目名称、项目编号/招标编号、招标人、招标代理、预算/最高限价、服务期限、服务地点、投标保证金、开标方式、评标方法等基础信息。

### 模块二：关键时间节点（必须放在项目基本信息之后）

提取所有时间相关的关键节点：招标公告发布时间、招标文件获取截止时间、答疑/澄清截止时间、投标截止时间（⚠️ 最关键）、开标时间、保证金缴纳截止时间（注意是否要求提前 N 个工作日到账）、有效期要求。

输出用倒计时时间轴，从当前日期（2026-06-23）开始计算剩余天数。关键时间节点必须作为报告第二章，放在项目基本信息速览之后。

### 模块三：废标条款（⚠️ 最高优先级）

找出所有一旦违反就导致投标无效的条款。

强制搜索关键词：
- 符号类：★ * ★★
- 拒绝词组：拒绝 / 无效 / 否决 / 不予通过 / 作废 / 废标 / 取消中标资格
- 强制性表述：应 / 必须 / 不得 / 严禁 / 禁止 / 未...视为不响应 / 未...按废标处理
- 审查相关：资格审查 / 符合性审查 / 初步评审 / 形式审查

重点扫描投标人须知、资格审查部分、评标办法章节。

对每条废标条款，必须记录：
- 招标文件页码（强制，必须精确到页）
- 原文引用（强制，完整句子，用引号包裹，不可改写或概括）
- 合规要求（需要做什么才能满足）
- 风险等级标注：致命（直接废标）/ 严重（扣大分）/ 需注意

特别关注带 ★ 或 * 标记的条款——这些通常是"一条不满足即废标"的星号条款。

### 模块四：评分标准及权重

完整还原评标打分表。找到「评标办法」「评分标准」「评审标准」等章节。

提取完整的评分结构：技术/商务/价格各部分权重和满分、每个评分项的序号/名称/满分/评分标准摘要/得分要点。

对于每个评分项，写出得分要点——即评委怎么给分的具体标准。

### 模块五：我方得分测算（上传方案/资质/产品材料时必须输出）

如果用户上传了方案、公司资质、业绩、人员、证书、产品参数、功能截图、报价或技术响应等材料，必须把这些材料与评分标准逐项匹配，估算我方可能得分。

得分测算必须分为两档：
- 悲观得分：只计算材料中明确可证明、评委较难扣分的项目；证据不足、表述不清、需补材料的项目按低分或不得分处理。
- 乐观得分：在材料基本符合、评委认可解释空间的前提下估算可争取分值；必须标明假设条件。

输出要求：
- 总分：悲观总分 / 乐观总分 / 满分。
- 分项表格：评分项、满分、悲观得分、乐观得分、证据来源（来自哪个上传文件/哪类材料）、扣分风险、补强动作。
- 不得凭空满分。若用户未上传可用于测算的方案/资质/产品材料，明确写“未上传我方材料，无法可靠估算我方得分”，但仍要列出需要补充的材料清单。

### 模块六：资质门槛

列出投标必须具备的所有资格条件。

提取范围：企业资质要求（营业执照范围、行业资质证书、等级要求）、财务要求（注册资本、资产负债率、审计报告年限）、业绩要求（类似项目业绩数量、金额门槛、时间范围）、人员要求（项目经理资质、团队配置、社保要求）、信用要求（无违法记录、信用中国截图等）、其他（ISO认证、安全生产许可证、特定行业许可等）。

对每项资质标注：必须具备（硬性门槛）/ 需要确认（可能有歧义）、提供证明材料的格式要求。

### 模块七：价格限制

提取：预算金额/最高限价、最低限价（如有）、报价方式（固定总价/固定单价/费率等）、保证金金额及缴纳方式、价格分计算公式（低价优先法/综合评分法/NLP法）、是否包含暂列金/暂估价。

特别关注价格分的计算逻辑——这直接影响报价策略。

### 模块八：加分机会、装订封装与投标执行

把所有能多拿分的地方全部找出来，并给出满分所需材料清单。

重点扫描方向：类似合同业绩、人员资历、客户评价、认证体系（ISO/CMMI）、知识产权（专利/软著）、技术亮点。

对每个加分项，必须写明：招标文件页码（强制）、分值、评分细则（原文引用，完整句子，不可改写或概括）、满分条件、需要提供的材料。
同时提取装订与封装要求：正本/副本份数、装订方式（胶装/线装/骑缝钉）、封装要求（密封章位置、封条格式）、电子版介质要求（U盘/光盘/光盘+U盘）、封套上的标识信息、签字盖章要求（哪些地方盖公章、法人章、骑缝章）。

## 投标文件框架生成规则

根据招标文件中的"投标文件组成""投标文件格式"章节生成投标文件框架。

核心原则：
- 只输出标题和子标题层级结构，不包含任何说明性文字、注释、预估页数或评分标注
- 用户拿到后可以直接全选→复制→粘贴进 Word 文档作为骨架
- 如果招标文件有明确要求的章节结构，严格遵照招标文件的目录结构，但去掉"附件X"编号前缀，只保留实际标题名称
- 目录结构中不要带"附件"二字
- 纯标题：不写括号里的补充说明、不写注释、不留空行占位
- 层级清晰：用四级标题即可
- 编号规范：使用中文数字（一、二、三）+ 阿拉伯数字（1.1/1.1.1）的层级编号方式

## 检查清单（Checklist）生成规则

生成投标完成后自查用的 Checklist，必须覆盖六个维度：
A. 废标条款核查（每条都必须勾选）
B. 内容完整性
C. 数据一致性（全文报价金额一致、公司名称统一、项目名称编号正确）
D. 格式规范（字体、行距、页边距、目录页码、页眉页脚）
E. 签章完整性（投标函、报价表、授权书、副本骑缝章、封套密封章）
F. 封装要求（正本份数、副本份数、电子版介质、装订顺序、密封、封套标识）

每个检查项用 checkbox 可点击打勾。

## 素材准备 TodoList 生成规则

生成素材收集任务清单，按三色分类：
- 🔴 紧急（影响投标资格）：营业执照、资质证书、审计报告、保证金、项目经理证书
- 🟡 重要（影响得分）：类似业绩合同、中标通知书、验收证明、团队人员证书
- 🟢 锦上添花（加分材料）：ISO认证、客户表扬信、专利软著、获奖证明

## 输出格式

输出一个完整的 HTML 文档（从 <!doctype html> 开始），不要使用 Markdown 代码围栏，不要输出 ```html 或 ```。包含以下 13 章，顺序必须一致：

1. 项目基本信息速览（项目名称、编号、招标人、招标代理、预算、截止时间等）
2. ⏰ 关键时间节点（必须紧跟基本信息）
3. 📊 评分标准详解
4. 🧮 我方得分测算（悲观/乐观）
5. ⚠️ 废标条款清单（按风险等级排列）
6. 📋 资质门槛清单
7. 💰 价格限制与报价策略参考
8. 🎯 加分机会与满分材料清单
9. 📦 装订封装要求
10. 📑 投标文件框架（可直接粘贴到 Word）
11. ✅ 投标文件自检 Checklist
12. 📝 素材准备 TodoList
13. 💡 综合分析与建议

## HTML 报告设计要求

报告应包含以下交互式特性：
1. 评分可视化：评分大类和分项用彩色条形图/进度条呈现，并突出“悲观得分 / 乐观得分 / 满分”的对比
2. 废标条款卡片：每条废标条款用醒目的红色/橙色卡片展示，风险等级用颜色区分
3. ★号条款高亮：星号否决条款用红色醒目卡片单独突出，标注"⚠️ 星号否决条款"
4. 时间轴可视化：关键时间节点用时间轴+倒计时天数+紧急标识展示
5. 可折叠章节：价格评分细则、投标文件目录等较长的内容可展开/收起
6. 交互式 Checklist：每个检查项可点击打勾标记完成状态
7. TodoList 分类卡片：紧急/重要/加分用三色分区卡片展示
8. 响应式设计：适配不同屏幕宽度，内容不会溢出

## HTML 样式参考

整体风格：专业商务风格，简洁清爽。以下样式可作为基础：
- body: font-family -apple-system, "PingFang SC", "Microsoft YaHei", sans-serif; background #f5f6fa; color #333
- 章节卡片: background white; border-radius 12px; padding 24px; margin 16px 0; box-shadow 0 2px 8px rgba(0,0,0,0.06)
- 危险卡片: background #fff5f5; border-left 4px solid #e74c3c
- 警告卡片: background #fff8e1; border-left 4px solid #ff9800
- 星号条款卡片: background #ffebee; border 2px solid #e74c3c
- 评分进度条: height 24px; border-radius 4px; background linear-gradient(90deg, #4caf50 var(--pct), #e0e0e0 var(--pct))
- 时间轴项: position relative; padding-left 32px; border-left 2px solid #2196f3
- 检查项: cursor pointer; padding 8px 12px; border-radius 8px; transition background 0.2s
- 可折叠: details/summary 标签

所有 CSS 内联在 <style> 标签中，不要引用外部资源。

## 重要规则

1. 宁可多提不可遗漏：对于不确定是否属于废标条款的内容，倾向于列入并标注"需确认"，而不是漏掉
2. 废标条款和加分机会必须标注招标文件原文页码并完整引用原文（用引号包裹，不得改写或概括），缺一不可。其他模块也遵循"尽量引用原文+标注页码"原则
3. 区分事实与推断：明确标注哪些是招标文件原文写的，哪些是基于经验的推断和建议
4. 投标文件框架必须是纯标题骨架：不包含任何说明、注释、占位符或预估信息
5. 报告末尾加免责声明："本报告由AI辅助生成，所有关键条款请以招标文件原文为准，建议人工复核后使用。"

## 摘要数据（必须包含）

在 HTML 的 <head> 中嵌入以下标签，供程序解析：

<script type="application/json" id="report-summary">
{{"project_name":"项目名称","project_no":"项目编号","buyer":"招标人","agent":"招标代理","budget":"预算/限价","deadline":"投标截止时间","risk_count":0,"score_total":0}}
</script>

其中 risk_count 为废标条款数量，score_total 为评分总分。请根据实际分析结果填写。

请直接输出完整 HTML，不要输出任何其他内容。不要用 Markdown 代码块包裹 HTML。"""


def call_llm(prompt: str) -> str:
    if not API_KEY:
        raise RuntimeError("MOONSHOT_API_KEY 环境变量未设置，无法调用大模型。")
    client = OpenAI(api_key=API_KEY, base_url=BASE_URL)
    model_name = MODEL.strip()
    # Kimi/Moonshot 部分模型只接受 temperature=1。
    temperature = 1 if model_name.lower().startswith("kimi-") else 0.3
    kwargs = {
        "model": model_name,
        "messages": [{"role": "user", "content": prompt}],
        "temperature": temperature,
        "stream": False,
    }
    try:
        response = client.chat.completions.create(**kwargs)
    except BadRequestError as exc:
        if "invalid temperature" not in str(exc).lower() or temperature == 1:
            raise
        kwargs["temperature"] = 1
        response = client.chat.completions.create(**kwargs)
    return response.choices[0].message.content or ""


def extract_summary(html: str) -> dict:
    match = re.search(
        r'<script[^>]*id="report-summary"[^>]*>(.*?)</script>',
        html,
        re.DOTALL,
    )
    if not match:
        return {}
    try:
        return json.loads(match.group(1).strip())
    except json.JSONDecodeError:
        return {}


def ensure_html(raw: str) -> str:
    raw = raw.strip()
    fence_match = re.match(r"^```(?:html)?\s*(.*?)\s*```$", raw, re.IGNORECASE | re.DOTALL)
    if fence_match:
        raw = fence_match.group(1)
    raw = sanitize_html(raw)
    raw = raw.strip()
    if raw.startswith("<!doctype") or raw.startswith("<!DOCTYPE"):
        return raw
    start = raw.lower().find("<html")
    if start != -1:
        return "<!doctype html>\n" + raw[start:]

    # 如果 LLM 输出的是纯文本或 Markdown，包装成基础 HTML
    return f"<!doctype html><html lang=\"zh-CN\"><head><meta charset=\"utf-8\"><style>body{{font-family:-apple-system,'PingFang SC','Microsoft YaHei',sans-serif;max-width:960px;margin:20px auto;padding:0 20px;line-height:1.65;color:#333}}h2{{border-bottom:2px solid #1a73e8;padding-bottom:6px;margin-top:32px}}.risk{{background:#fff5f5;border-left:4px solid #d93025;padding:12px 16px;margin:10px 0;border-radius:6px}}pre,code{{background:#f5f5f5;padding:2px 5px;border-radius:3px;font-size:.9em}}</style></head><body>{raw}</body></html>"


def sanitize_html(html_content: str) -> str:
    """使用 bleach 清洗 HTML，移除 XSS 攻击向量。

    只允许安全的标签、属性和 CSS 属性。移除 script、事件处理器、危险 URL 等。
    """
    allowed_tags = {
        "html", "head", "body", "meta", "title", "link", "style",
        "h1", "h2", "h3", "h4", "h5", "h6",
        "p", "br", "hr", "blockquote", "pre", "code",
        "ul", "ol", "li", "dl", "dt", "dd",
        "table", "thead", "tbody", "tfoot", "tr", "th", "td", "caption", "colgroup", "col",
        "div", "span", "section", "article", "header", "footer", "nav", "aside", "main",
        "a", "strong", "b", "em", "i", "u", "s", "small", "sub", "sup", "mark",
        "img", "figure", "figcaption",
        "details", "summary",
        "label", "input", "button", "form",
        "script",  # 仅允许 type="application/json" 用于数据嵌入
        "style",
    }

    allowed_attrs = {
        "*": ["class", "id", "style", "lang", "title", "dir"],
        "a": ["href", "target", "rel"],
        "img": ["src", "alt", "width", "height", "loading"],
        "meta": ["charset", "name", "content"],
        "link": ["rel", "href", "type"],
        "td": ["colspan", "rowspan"],
        "th": ["colspan", "rowspan", "scope"],
        "input": ["type", "checked", "disabled"],
        "label": ["for"],
        "details": ["open"],
        "summary": [],
        "script": ["type", "id"],  # 仅允许 application/json 类型
        "col": ["span"],
        "colgroup": ["span"],
    }

    allowed_protocols = ["http", "https", "mailto"]

    # CSS 白名单 — 允许布局和样式，禁止 expression() 和行为绑定
    allowed_css = [
        "color", "background-color", "background", "border", "border-color",
        "border-radius", "border-left", "border-right", "border-top", "border-bottom",
        "padding", "padding-left", "padding-right", "padding-top", "padding-bottom",
        "margin", "margin-left", "margin-right", "margin-top", "margin-bottom",
        "width", "max-width", "min-width", "height", "max-height", "min-height",
        "font-family", "font-size", "font-weight", "font-style",
        "line-height", "text-align", "text-decoration", "text-transform",
        "letter-spacing", "word-spacing", "white-space",
        "display", "position", "top", "right", "bottom", "left",
        "overflow", "overflow-x", "overflow-y",
        "box-shadow", "opacity", "visibility",
        "cursor", "grid-template-columns", "grid-template-rows", "grid-gap", "gap",
        "flex", "flex-direction", "flex-wrap", "justify-content", "align-items",
        "align-self", "order",
        "z-index", "transition", "transform",
        "list-style", "content",
        "scroll-behavior", "box-sizing",
        "background-image", "background-size", "background-position", "background-repeat",
        "border-collapse", "table-layout", "vertical-align",
        "word-break", "word-wrap", "overflow-wrap", "text-overflow",
        "outline", "resize",
        "inset", "place-items", "object-fit",
        "fill", "stroke", "stroke-width",
        # CSS 变量和渐变
        "linear-gradient", "radial-gradient",
    ]

    def _strip_unsafe_script(match: re.Match[str]) -> str:
        attrs = match.group(1)
        content = match.group(2)
        is_json = re.search(r'\btype\s*=\s*["\']application/json["\']', attrs, re.IGNORECASE)
        is_summary = re.search(r'\bid\s*=\s*["\']report-summary["\']', attrs, re.IGNORECASE)
        if is_json and is_summary:
            return f"<script{attrs}>{content}</script>"
        return ""

    html_content = re.sub(
        r"<script\b([^>]*)>(.*?)</script>",
        _strip_unsafe_script,
        html_content,
        flags=re.IGNORECASE | re.DOTALL,
    )

    css_sanitizer = CSSSanitizer(
        allowed_css_properties=[
            prop for prop in allowed_css
            if prop not in {"fill", "stroke", "stroke-width", "linear-gradient", "radial-gradient"}
        ],
        allowed_svg_properties=["fill", "stroke", "stroke-width"],
    )

    # 使用 bleach.Cleaner 做更精细的清洗（CSS 过滤 + script 预处理）
    cleaner = bleach.Cleaner(
        tags=allowed_tags,
        attributes=allowed_attrs,
        protocols=allowed_protocols,
        css_sanitizer=css_sanitizer,
        strip=True,
        strip_comments=True,
    )
    cleaned = cleaner.clean(html_content)

    # 移除内联事件处理器（onclick, onload 等）—— bleach 6.x 默认应该做了
    # 但为保险起见再跑一次正则清除
    event_attrs = re.compile(
        r'\bon[a-z]+\s*=\s*["\'][^"\']*["\']',
        re.IGNORECASE,
    )
    cleaned = event_attrs.sub("", cleaned)

    # 移除 javascript: 伪协议（bleach 应该已处理，冗余保护）
    cleaned = re.sub(
        r'(?i)href\s*=\s*["\']javascript:',
        'href="#" data-removed="javascript-uri"',
        cleaned,
    )

    return cleaned


def ensure_html_closed(html: str) -> str:
    """补全可能因 token 截断而缺失的闭合标签"""
    tags_to_close = []
    for tag in ["li", "ul", "ol", "div", "p", "section", "article",
                 "details", "blockquote", "table", "tbody", "tr", "td", "th",
                 "span", "strong", "em", "a", "h1", "h2", "h3", "h4", "h5", "h6",
                 "script", "style", "head"]:
        open_pattern = rf"<{tag}(?:\s[^>]*)?>"
        close_pattern = rf"</{tag}>"
        opens = len(re.findall(open_pattern, html, re.IGNORECASE))
        closes = len(re.findall(close_pattern, html, re.IGNORECASE))
        deficit = opens - closes
        if deficit > 0:
            tags_to_close.extend([f"</{tag}>"] * deficit)
    # 确保 body 和 html 闭合
    if "</body>" not in html:
        tags_to_close.append("</body>")
    if "</html>" not in html:
        tags_to_close.append("</html>")
    if tags_to_close:
        html = html + "\n" + "".join(reversed(tags_to_close)) + "\n<!-- [note: output may have been truncated by model token limit] -->"
    return html


def main() -> None:
    tender = Path(os.environ["BIDDING_ANALYST_TENDER"])
    sca = Path(os.environ["BIDDING_ANALYST_SCA"]) if os.environ.get("BIDDING_ANALYST_SCA") else None
    out = Path(os.environ["BIDDING_ANALYST_OUT"])

    if not tender.exists():
        print(f"招标文件不存在：{tender}", file=sys.stderr)
        sys.exit(1)

    tender_text = extract_text(tender)
    if not tender_text.strip():
        print(f"无法从招标文件提取文本：{tender}", file=sys.stderr)
        sys.exit(1)

    solution_text = None
    if sca and sca.exists():
        solution_text = extract_text(sca)

    prompt = build_prompt(tender_text, solution_text)
    print(f"[core] 调用大模型 {MODEL}，招标文件 {len(tender_text)} 字...", file=sys.stderr)

    raw_html = call_llm(prompt)
    html = ensure_html(raw_html)
    html = ensure_html_closed(html)

    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_text(html, encoding="utf-8")

    summary = extract_summary(html)
    summary["engine"] = "llm-moonshot"
    summary["report_path"] = str(out)
    summary["generated_at"] = __import__("datetime").datetime.now().isoformat(timespec="seconds")

    print(json.dumps(summary, ensure_ascii=False))


if __name__ == "__main__":
    main()
